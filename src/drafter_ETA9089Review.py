from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import config

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

class ETA9089drafter:

    def __init__(self):
        self.length = len(config.information)
        self.document = Document()
        self.attorneys = {'ATT': 'AttorneyName@Org.com'}
        #### Declaring the styles of the document prior to writing anything down
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Arial'
        self.font.size = Pt(10)

    def mainloop(self):
    #### main loop function to write the document
        for i in range(0, self.length):
            attorney = self.attorneys.get(config.information[i][3])
            self.document.add_heading("ETA 9089 for Employee Review, Generics", 0)
            ####Heading, Send to:, CC:, and subject line drafted here
            p = self.document.add_paragraph()
            p.add_run("Send To: ").bold = True
            p.add_run("{FNemail}\n\n".format(FNemail=config.information[i][4]))
            p.add_run("CC: ").bold = True
            p.add_run("{att}\n\n".format(att=str(attorney)))
            p.add_run("Subject: ").bold = True
            p.add_run("{LAST}, {First} - PERM C# {case} - Employee Review ETA 9089 [Action Required]\n\n".format(LAST=config.information[i][0],
                                                                                                            First=config.information[i][1],
                                                                                                            case=str(config.information[i][2])))
            ####Email Body
            ####Banner
            self.document.add_picture('images/ETA9089Review.png', width=Inches(5.5))
            ### Paragraph 1 "p1" should be the same for all different options.
            p1 = self.document.add_paragraph()
            p1.style = self.document.styles['Normal']
            p1.add_run('Classification: Confidential\n\n')
            p1.add_run("Dear {first},\n\n".format(first = config.information[i][1]))
            p1.add_run("Company name is actively conducting a Labor Market Test for the PERM application and has "
                       "authorized us to prepare Form ETA 9089 (the final PERM application form) on your behalf. "
                       "Attached please find a draft version of the Form ETA 9089. "
                       "Before Company submits an application to the U.S. Department of Labor (DOL), you must review the "
                       "portions of the form that contain your personal information as well as information on your education and prior work experience. "
                       "This form has been prepared based on the information you provided in your Employee Draft Review "
                       "and any verification letters you have provided. \n\n"
                       "Before Company submits an application to the U.S. Department of Labor (DOL), you must review the "
                       "portions of the form that contain your personal information as well as information on your education "
                       "and prior work experience. This form has been prepared based on the information you provided in your "
                       "Employee Draft Review and any verification letters you have provided. ")
            ##### Need to do image
            self.document.add_picture('images/needtodo.png', width = Inches(5.5))
            ##### Paragraph 2 plus table
            ###
            p2 = self.document.add_paragraph()
            p2.style = self.document.styles['Normal']
            p2.add_run("Immediately\n").bold = True
            p2.add_run("Please carefully review the attached sections of the ETA 9089. It is essential that the information"
                       " on the form is accurate. Any inaccuracies in the information provided can result in the denial "
                       "of the PERM Labor Certification or the inability to utilize an approved PERM application for the "
                       "next steps in the green card process. Every section is important. ")
            ########### Table
            table = self.document.add_table(rows=3, cols=2)
            table.style = 'Table Grid'

            column = table.columns[0]
            column.cells[0].text = 'Confirmations'
            column.cells[1].text = "I have reviewed Section J and can confirm that the information in Section J is correct."
            column.cells[2].text = "I have reviewed Section K and can confirm that the job history in section K is correct."

            column = table.columns[1]
            column.cells[0].text = "Your Response: (Yes or No. If 'No', please explain)"

            make_rows_bold(table.rows[0])
            ###############
            p3 = self.document.add_paragraph()
            p3.style = self.document.styles['Normal']
            p3.add_run("\n\n\n\nIf there is ")
            p3.add_run("any ").bold = True
            p3.add_run("information on this form that requires correction, ")
            p3.add_run("please let us know immediately.").bold = True
            p3.add_run("The application, once submitted to the DOL, cannot be changed; therefore, you should review its contents very carefully.\n\n")

            p3.add_run("Section J.9 and J.10 can be left blank. However, if you have a new I-94 document, please upload it to the Connect "
                       "portal or email it to Company@organization.com, so we can update your file.\n\nPlease note: The attachment "
                       "is password protected, for your security. We will send the password in a separate email to follow shortly.\n")
            #####
            ################
            ### Next steps image
            self.document.add_picture("images/next_steps.png", width = Inches(5.5))
            ### Paragraph 3
            p4 = self.document.add_paragraph()
            p4.style = self.document.styles['Normal']
            p4.add_run("1. 	Once the labor market test is complete, there is a mandatory cool-down period to allow your manager "
                       "time to review resumes from U.S. applicants.  The duration of the cool-down period may not be shorter than 30 days, "
                       "but it may be longer depending on the volume of responses to the advertisements and how quickly they can be reviewed.\n\n ")
            p4.add_run("2. 	Provided there are no qualified U.S. applicants, we will then file Company’s PERM application. The date the PERM application "
                       "is filed with the U.S. Department of Labor (DOL) will become your priority date. Based on current DOL processing times, "
                       "it will take approximately 4-6 months to receive a response.\n\n ")
            p4.add_run("3.	Once we receive a response from DOL, we will reach out to you with updates. After the PERM application "
                       "is certified (approved), our legal team will automatically reach out to you to initiate the second stage of the "
                       "permanent residence process, the I-140 petition.\n")
            ####
            #### Flowchart Image
            self.document.add_picture("images/flowchart.png", width = Inches(5.5))
            #### Closing Paragraph
            p5 = self.document.add_paragraph()
            p5.add_run("Please let me know if you have any questions.\n\nThank you and I look forward to receiving your documents.\n\n"
                       "Kind Regards,")

            ###### Follow up email with the password
            ####Heading, Send to:, CC:, and subject line drafted here
            p = self.document.add_paragraph()
            p.add_run("Send To: ").bold = True
            p.add_run("{FNemail}\n\n".format(FNemail=config.information[i][4]))
            p.add_run("Subject: ").bold = True
            p.add_run("{LAST}, {First} - PERM C# {case} - Employee Review ETA 9089 Document Password".format(
                LAST=config.information[i][0],
                First=config.information[i][1],
                case=str(config.information[i][2])))
            ####Email Body
            p1 = self.document.add_paragraph()
            p1.add_run("Hello,\n\n")
            p1.add_run("The document is password protected and attached to the previous email. The password is as follows: "
                       "Immigration20!\n\nThat is a capital ‘i’ at the front and an exclamation point at the end. "
                       "Please let me know if you cannot access the document or if you have any questions. Thank you. "
                       "\n\nKind regards,"
                       "\n\nJerome")
            self.document.add_page_break()
        self.document.save('output/ETA 9089 for Employee Review.docx')
    ############

##main function of the class
def ETA9089draft():
    doc = ETA9089drafter()
    doc.mainloop()