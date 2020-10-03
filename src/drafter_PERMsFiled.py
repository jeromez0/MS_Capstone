from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import config


class PERMsfiled:

    def __init__(self):
        self.length = len(config.information)
        self.document = Document()
        self.information = config.information
        #### Declaring the styles of the document prior to writing anything down
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Arial'
        self.font.size = Pt(10)

        self.attorneys = {'ATT': 'AttorneyName@Org.com'}

    def createdocument(self):
        for i in range(0, self.length):
            attorney = self.attorneys.get(self.information[i][3])
            p = self.document.add_paragraph()

            p.add_run('Send To: ').bold = True
            p.add_run("{FNemail}\n\n".format(FNemail = self.information[i][4]))
            p.add_run('CC: ').bold = True
            p.add_run("{att}, {manageremail}\n\n".format(att = attorney,
                                                         manageremail = self.information[i][5]))
            p.add_run('Subject Line: ').bold = True
            p.add_run('{Last}, {First} - Company Name - PERM Filed - Case #{number}'.format( Last = self.information[i][0],
                                                                                                    First = self.information[i][1],
                                                                                                    number = str(self.information[i][2])))
            #### Email Body
            #### banner
            self.document.add_picture('images/filing.jpg', width=Inches(5.5))

            ### Paragraph 1
            p = self.document.add_paragraph()
            p.style = self.document.styles['Normal']
            p.add_run('Dear {first},\n\n'.format(first = self.information[i][1]))
            p.add_run("We confirm that we have filed Company Name's PERM Application under the ")

            if self.information[6] == 'A':
                p.add_run("EB-2 ")
            else:
                p.add_run("EB-3 ")

            p.add_run("preference category on your behalf. We will contact you once we have an update. \n\n")
            p.add_run('Processing of the application is expected to take approximately 4-5 months.')
            self.document.add_picture('images/next_steps.png', width=Inches(5.5), height=Inches(0.5))

            ### Paragraph 2
            p = self.document.add_paragraph()
            p.style = self.document.styles['Normal']
            p.add_run(
                '\n1. The date that we filed the PERM application is your priority date. It will take approximately '
                '4-5 months before we receive a decision from the Department of Labor.\n\n')
            p.add_run('2. Once the PERM application is approved, we will be contacting you and requesting the following'
                      ' current information/documentation required for the I-140 petition:')

            ### Bulleted list
            self.document.add_paragraph('Clear copy of most recent I-94', style='List Bullet')
            self.document.add_paragraph('Your current address', style='List Bullet')
            self.document.add_paragraph("A permanent foreign address (could be a relative's)", style='List Bullet')
            self.document.add_paragraph("Current work address", style='List Bullet')
            self.document.add_paragraph("Current salary", style='List Bullet')
            self.document.add_paragraph(
                "Information regarding immediate family including name of spouse and children; dates of birth; country of "
                "birth; and present address (if different than yours)\n", style='List Bullet')
            self.document.add_picture('images/wherearewe.jpg', width=Inches(5), height=Inches(0.5))

            ###Final paragraph
            p = self.document.add_paragraph()
            p.style = self.document.styles['Normal']
            p.add_run(
                'Please contact our office prior to making, amending, or confirming any travel plans or assignment start dates.\n\n')
            p.add_run('Kind Regards,')
        self.document.save('output/PERMsFiled.docx')

def mainfunction():
    PERM = PERMsfiled()
    PERM.createdocument()