import config
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from datetime import datetime

def make_columns_bold(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

class doc_drafter:
####class declaration JD&Reqs drafter. Includes all the attributes of the JD&Reqs document.
    def __init__(self):
        self.length = len(config.information)
        self.document = Document()

        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Calibri'
        self.font.size = Pt(10)

    def drafter(self):
        testing = ["ATTY/PARA: ", "FN NAME: ", "FNIV", "JOB CODE (JO#): ", "ETA ACCOUNT: ",  ###testing is the list that will store
                   "COB: ", "CASE #: ", "MAJOR (DATE): ", "SALARY RANGE: ", "ETA CASE #: "]  ### the different table values.
        ##### the for loop will create however many JD&Req documents are needed.
        for i in range(0, self.length):
            self.document = Document()
            self.style = self.document.styles['Normal']
            self.font = self.style.font
            self.font.name = 'Calibri'
            self.font.size = Pt(10)

        ########### First table with all the relevant FN information
            table = self.document.add_table(rows = 5, cols = 4)
            table.style = 'Table Grid'

            column = table.columns[0]
            column.cells[0].text = testing[0]
            column.cells[1].text = testing[1]
            column.cells[2].text = testing[2]
            column.cells[3].text = testing[3]
            column.cells[4].text = testing[4]

            column = table.columns[2]
            column.cells[0].text = testing[5]
            column.cells[1].text = testing[6]
            column.cells[2].text = testing[7]
            column.cells[3].text = testing[8]
            column.cells[4].text = testing[9]

            make_columns_bold(table.columns[0], table.columns[2])
        ##########
            ####This block inputs all the proper FN information.
            column = table.columns[1]
            column.cells[0].text = config.information[i][3] + " / JZ"
            column.cells[1].text = config.information[i][0] + ", " + config.information[i][1]
            try:
                column.cells[2].text = datetime.strftime(config.information[i][4], '%m/%d/%Y')
            except:
                column.cells[2].text = str(config.information[i][4])

            column = table.columns[3]
            column.cells[0].text = config.information[i][5]
            column.cells[1].text = str(config.information[i][2])
            column.cells[2].text = config.information[i][7]
            ####
        ###########
        ## Second table called Important Case Notes
        ###########
            p = self.document.add_paragraph()
            p.add_run("\n")
            table = self.document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0,0)
            cell.text = "                             IMPORTANT CASE NOTES:"
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(14)

            table = self.document.add_table(rows = 1, cols = 1)
            table.style = 'Table Grid'
            cell = table.cell(0,0)
            cell.text = "XX/XX/XXXX:   |"
            p = self.document.add_paragraph()

        #############
        #### JD & Requirements header table
        #############
            table = self.document.add_table(rows=1,cols = 1)
            table.style = 'Table Grid'
            cell = table.cell(0,0)
            cell.text = "                             JD & REQUIREMENTS:  "
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(14)

            table = self.document.add_table(rows = 1, cols = 4)
            table.style = 'Table Grid'
            cell = table.cell(0,0)
            cell.text = "Reviewed by:"
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)

            cell = table.cell(0,1)
            cell.text = ('Manager - via email')
            run = cell.paragraphs[0].runs[0]
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)

            cell = table.cell(0,2)
            cell.text = "Date:"
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)
        #############

        #### JD & Requirements Body Table
        #############
            table = self.document.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
        ################ Job Title
            cell.text = "\nJob Title:"
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run = cell.paragraphs[0]
            run.add_run(" Software Engineer")
        ###############  Job Description
            run.add_run("\n\nJob Description: ").bold = True
            run.add_run("Responsible for the definition, design, development, test, debugging, release, "
                    "enhancement or maintenance of networking software. ").bold = True

            run.add_run("Write design specifications and functional specifications. Design, develop and define new software "
                    "product features. Perform complex system level tests. Solve problems of more complex scope "
                    "involving multiple modules. Act as an individual contributor under limited supervision. "
                    "Exercise independent judgment within defined company standards. ")
            ##### This is a conditional. Depending on the requirements, an additional line may be needed.
            if config.information[i][9] == 'B':
                run.add_run("Apply advanced theoretical knowledge and skills to perform job duties.").bold = True
            elif config.information[i][9] == 'C1':
                run.add_run("Apply advanced theoretical knowledge and skills to perform job duties.").bold = True
            elif config.information[i][9] == 'C2':
                pass
            elif config.information[i][9] == 'A':
                pass

        ############### Job Site
            run.add_run("\n\nJob Site:").bold = True
            run.add_run(" 170 W Tasman Drive, San Jose, CA 95134\n\n")
            run.add_run("Salary: ").bold = True
            run.add_run("${salary}/year".format(salary = config.information[i][6]))
        ############### Education and Experience Requirements
            #### The requirements will vary depending on each option.
            run.add_run("\n\nEducational Requirements:").bold = True
            if config.information[i][9] == 'B':
                run.add_run(" Master's or foreign degree equivalent in Computer Science, Computer Engineering, Electrical "
                        "Engineering or related field.\n\n")
                run.add_run("Experience Requirements:").bold = True
                run.add_run(" Six (6) months of experience as a Software Engineer or in a related occupation.")

            if config.information[i][9] == 'C1':
                run.add_run(" Master's or foreign degree equivalent in Computer Science, Computer Engineering, Electrical "
                        "Engineering, or related field.\n\n")
                run.add_run("Experience Requirements:").bold = True
                run.add_run(" Three (3) years of experience as a Software Engineer or in a related occupation.")

            if config.information[i][9] == 'C2':
                run.add_run(" Bachelor's or foreign degree equivalent in Computer Science, Computer Engineering, Electrical "
                        "Engineering or related field.\n\n")
                run.add_run("Experience Requirements:").bold = True
                run.add_run(" Five (5) years of post-baccalaureate, progressive experience as a Software Engineer or "
                        "in a related occupation.")

            if config.information[i][9] == 'A':
                run.add_run(" Bachelor's or foreign degree equivalent in Computer Science, Computer Engineering, Electrical "
                        "Engineering, or related field.\n\n")
                run.add_run("Experience Requirements:").bold = True
                run.add_run(" Two (2) years of experience as a Software Engineer or in a related occupation.")
                ###
        ##############
            ### special skills requirement
            run.add_run("\n\nSpecial Skills:\n").bold = True
            run.add_run("* Network Protocols and/or Switching Techniques for networks.\n\n")
            ### other notes
            p = self.document.add_paragraph()
            p.add_run("Other Notes:")
            #####
            ####
            ####the document needs to be saved according to the option type as well as the FN's name.
            self.document.save("output/JD&Reqs/0. JD&Reqs Option {opt} ({LAST}, {First}).docx".format(opt = config.information[i][9],
                                                                                              LAST = config.information[i][0],
                                                                                             First = config.information[i][1]))

def maindraft():
    JD = doc_drafter()
    JD.drafter()
