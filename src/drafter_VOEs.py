from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from datetime import datetime
from datetime import timedelta
from docx.shared import RGBColor
import config

class createdocument:

    def __init__(self):
        self.length = len(config.information)
        self.document = Document()
        self.oneweekfromnow = datetime.now() + timedelta(days=7)
        self.oneweekfromnow = datetime.strftime(self.oneweekfromnow, '%m/%d/%Y')

        #### Declaring the styles of the document prior to writing anything down
        self.style = self.document.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Arial'
        self.font.size = Pt(10)

        self.attorneys = {'ATT': 'AttorneyName@Org.com'}

    ###################### This function determines each option and will output the proper requirements per option.

    def mainloop(self):
    #### main loop function to write the document
        for i in range(0, self.length):
            attorney = self.attorneys.get(config.information[i][3])
            self.document.add_heading("VOE Email's, Generics", 0)
            ####Heading, Send to:, CC:, and subject line drafted here
            p = self.document.add_paragraph()
            p.add_run("Send To: ").bold = True
            p.add_run("{FNemail}\n\n".format(FNemail = config.information[i][4]))
            p.add_run("CC: ").bold = True
            p.add_run("{manageremail}, {att}\n\n".format(manageremail = config.information[i][5],
                                                         att = str(attorney)))
            p.add_run("Subject: ").bold = True
            p.add_run("{LAST}, {First} - PERM C# {case} - Option {opt} Verification of Employment Instructions"
                      "[ACTION REQUIRED BY {date}]\n\n".format(LAST = config.information[i][0],
                                                               First = config.information[i][1],
                                                               case = str(config.information[i][2]),
                                                                opt = config.information[i][6],
                                                                date = self.oneweekfromnow))
        ####Email Body
        ####Banner
            self.document.add_picture('images/header.png', width=Inches(5.5))
        ### Paragraph 1 "p1" should be the same for all different options.
            p1 = self.document.add_paragraph()
            p1.style = self.document.styles['Normal']
            p1.add_run('Classification: Confidential\n\n')
            p1.add_run("Dear " + config.information[i][1] + ",\n\n")
            p1.add_run("Your manager has confirmed the education, experience, and special skills requirements for the PERM "
                   "application as follows:\n\n")
            p12 = p1.add_run("Job Title: Software Engineer\nWorksite: San Jose, CA\n\n")
            p12.bold = True
            font = p12.font
            font.color.rgb = RGBColor(0, 0, 255)
            p11 = p1.add_run("Education and Experience:")
            p11.bold = True
            p11.underline = True
            font = p11.font
            font.color.rgb = RGBColor(0, 0, 255)
        ### Paragraph 2 "p2" will be different depending on case option. We will use an if conditional gate.

            p2 = self.document.add_paragraph()
            p2.style = self.document.styles['Normal']

            #### If conditional for option A
            if config.information[i][6] == 'A':
                p20 = p2.add_run("* Bachelor's or foreign degree equivalent\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p2.add_run("* Two (2) years of full-time experience\n\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p21 = p2.add_run("SPECIAL SKILLS\n\n")
                p21.bold = True
                p21.underline = True
                font = p21.font
                font.color.rgb = RGBColor(0, 0, 255)
                p22 = p2.add_run("* Network Protocols and/or Switching Techniques for networks\n\n")
                font = p22.font
                font.color.rgb = RGBColor(0, 0, 255)
                p2.add_run(
                    "The job description and minimum requirements were developed based on the information you and "
                    "your manager provided. Based upon these requirements, the I-140 (the next step in the green "
                    "card process) should be adjudicated in the ")
                p2.add_run("EB-3 category.").bold = True

            #### If conditional for option B
            if config.information[i][6] == 'B':
                p20 = p2.add_run("* Master's or foreign degree equivalent\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p20 = p2.add_run("* Six (6) months of full-time experience\n\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p21 = p2.add_run("SPECIAL SKILLS\n\n")
                p21.bold = True
                p21.underline = True
                font = p21.font
                font.color.rgb = RGBColor(0, 0, 255)
                p22 = p2.add_run("* Network Protocols and/or Switching Techniques for networks\n\n")
                font = p22.font
                font.color.rgb = RGBColor(0, 0, 255)
                p2.add_run(
                    "The job description and minimum requirements were developed based on the information you and "
                    "your manager provided. Based upon these requirements, the I-140 (the next step in the green "
                    "card process) should be adjudicated in the ")
                p2.add_run("EB-2 category.").bold = True

            #### If conditional for option C1
            if config.information[i][6] == 'C1':
                p20 = p2.add_run("* Master's or foreign degree equivalent\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p20 = p2.add_run("* Three (3) years of full-time experience\n\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p21 = p2.add_run("SPECIAL SKILLS\n\n")
                p21.bold = True
                p21.underline = True
                font = p21.font
                font.color.rgb = RGBColor(0, 0, 255)
                p22 = p2.add_run("* Network Protocols and/or Switching Techniques for networks\n\n")
                font = p22.font
                font.color.rgb = RGBColor(0, 0, 255)
                p2.add_run(
                    "The job description and minimum requirements were developed based on the information you and "
                    "your manager provided. Based upon these requirements, the I-140 (the next step in the green "
                    "card process) should be adjudicated in the ")
                p2.add_run("EB-2 category.").bold = True

            #### If conditional for option C2
            if config.information[i][6] == 'C2':
                p20 = p2.add_run("* Bachelor's or foreign degree equivalent\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p20 = p2.add_run("* Five (5) years of post-baccalaureate, progressive full-time experience**\n\n")
                p20.bold = True
                font = p20.font
                font.color.rgb = RGBColor(0, 0, 255)
                p21 = p2.add_run("SPECIAL SKILLS\n\n")
                p21.bold = True
                p21.underline = True
                font = p21.font
                font.color.rgb = RGBColor(0, 0, 255)
                p22 = p2.add_run("* Network Protocols and/or Switching Techniques for networks\n\n")
                font = p22.font
                font.color.rgb = RGBColor(0, 0, 255)
                p2.add_run(
                    "The job description and minimum requirements were developed based on the information you and "
                    "your manager provided. Based upon these requirements, the I-140 (the next step in the green "
                    "card process) should be adjudicated in the ")
                p2.add_run("EB-2 category.\n\n")
                p23 = p2.add_run(
                    "**You will be required to demonstrate your post-baccalaureate experience was 'progressive' in nature.")
                p23.bold = True
                font = p23.font
                font.color.rgb = RGBColor(72, 61, 139)
                p24 = p2.add_run(
                    " This means that you “progressed” over time in your employment, taking on more responsibilities "
                    "in terms of job duties or supervision. Progressive experience might be shown through an "
                    "increase in job responsibilities within the same position, grade/level increases or "
                    "promotions, or progression to a job title with increasing responsibilities within the same "
                    "company or from one employment to another (e.g. Software Engineer >> Sr. Software Engineer "
                    ">> Tech Lead >> Management). ")
                font = p24.font
                font.color.rgb = RGBColor(72, 61, 139)
            #################################################

        ####Another image added
            self.document.add_picture("images/needtodo.png", width=Inches(5.5))

        ### Paragraph 3 = "p3"
            p3 = self.document.add_paragraph()
            p30 = p3.add_run("1. ASAP")
            p30.bold = True
            p30.underline = True
            font = p30.font
            font.color.rgb = RGBColor(30, 144, 255)
            p3.add_run(
            "\n\nAt this time, please confirm that you will be able to obtain VOE and skills letters to document "
            "the minimum requirements and special skills for the proposed position. If you will be unable to "
            "verify any of the special skills and/or the minimum experience requirements, you must"
            " inform us immediately.\n\n").bold = True

            p31 = p3.add_run("2. Prior to Filing the PERM Application\n")
            p31.bold = True
            p31.underline = True
            font = p31.font
            font.color.rgb = RGBColor(30, 144, 255)
            p3.add_run("Provide evidence verifying that you met all experience and/or skills requirements (listed "
                   "above prior to joining Company U.S.\n\n").bold = True
            p3.add_run("Verification of Employment (VOE) letters should include the following:")
        ####

        ####Section 4 of the email, bulleted lists
            self.document.add_paragraph("Dates of Employment", style='List Bullet')
            self.document.add_paragraph("Hours worked per week (or that you were employed full-time)", style='List Bullet')
            self.document.add_paragraph("Job Title", style='List Bullet')
            self.document.add_paragraph("Brief description of job duties", style='List Bullet')
            self.document.add_paragraph("Verification of any skills requirements you plan to show through the employer",
                               style='List Bullet')
            self.document.add_paragraph("Signed by a supervisor/manager or company official (such as H.R. on official "
                               "company letterhead", style='List Bullet')
            p4 = self.document.add_paragraph()
            p4.add_run('You may also verify skills requirements through your education with a letter from a professor '
                   'signed on university letterhead.\n\nWe have attached sample verification letter templates '
                   'with this email along with a “Verification of Employment (VOE) Frequently Asked Questions (FAQs)” '
                   'document, for your reference.\n\n\n\n\n')
            p4.add_run("Please Note: ")
            p41 = p4.add_run("Final signed letters will be required before PERM Applications are filed.")
            p41.bold = True
            font = p41.font
            font.color.rgb = RGBColor(255, 0, 0)
        ####
        ####Another image
            self.document.add_picture("images/next_steps.png", width=Inches(5.5), height=Inches(0.5))

        ####Paragraph/Section 5
            p5 = self.document.add_paragraph()
            p5.add_run("1.	Once the we receive confirmation that you can obtain the VOE and skills letters, we will file "
                   "the request for a Prevailing Wage Determination with the U.S. Department of Labor (DOL).\n\n"
                   "2. 	Layoff analysis is performed, and if possible, we will move forward with labor market testing, "
                   "which typically takes approximately 3-4 months, if no issues arise. \n\n"
                   "3.	We will proceed to draft the Form ETA 9089 and send it to you for review. If no qualified U.S. worker "
                   "is found in the recruitment process, we can file Company’s PERM application (Form ETA 9089) with the DOL.\n\n")
        ####
        ####Another Image
            self.document.add_picture("images/flowchart.png", width=Inches(5.5), height=Inches(.75))

        #### Paragraph/Section 6 (Final Paragraph)
            p6 = self.document.add_paragraph()
            p6.add_run("If you have any questions, please refer to our VOE FAQs document attached to this email. "
                   "If you still have questions after consulting the FAQs, please respond directly to this email. "
                   "Thank you.\n\n"
                   "Kind regards,\n")
        #####
            self.document.add_page_break()
    ####Save document as such
        self.document.save('output/Automated VOE Emails (Generics).docx')

def mainfunc():
    VOE_Doc = createdocument()
    VOE_Doc.mainloop()